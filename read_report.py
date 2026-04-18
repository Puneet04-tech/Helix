from docx import Document

try:
    doc = Document('AI_Guardian_Master_Report.docx')
    print("=" * 80)
    print("AI GUARDIAN MASTER REPORT CONTENT")
    print("=" * 80)
    
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    
    print("\n\n" + "=" * 80)
    print("TABLES IN DOCUMENT")
    print("=" * 80 + "\n")
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTable {table_idx + 1}:")
        for row_idx, row in enumerate(table.rows):
            print(" | ".join([cell.text.strip() for cell in row.cells]))
    
except Exception as e:
    print(f"Error reading document: {e}")
    import traceback
    traceback.print_exc()
