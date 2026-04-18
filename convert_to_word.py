#!/usr/bin/env python3
"""
Convert HELIX_MASTER_REPORT_COMPLIANCE_VERIFICATION.md to Word document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def read_markdown_file(filepath):
    """Read markdown file"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def add_heading_with_level(doc, text, level=1):
    """Add heading with appropriate level"""
    # Remove markdown markers
    text = text.lstrip('#').strip()
    try:
        doc.add_heading(text, level=level)
    except:
        doc.add_paragraph(text, style='Heading 1')

def add_table_from_markdown(doc, markdown_table):
    """Parse and add markdown table to document"""
    lines = markdown_table.strip().split('\n')
    if len(lines) < 3:
        return False
    
    # Parse header
    header = [cell.strip() for cell in lines[0].split('|')]
    header = [h for h in header if h]  # Remove empty
    
    if not header:
        return False
    
    # Parse rows
    rows = []
    for i in range(2, len(lines)):
        if '---' in lines[i]:
            continue
        cells = [cell.strip() for cell in lines[i].split('|')]
        cells = [c for c in cells if c]  # Remove empty
        if cells:
            rows.append(cells)
    
    if not rows:
        return False
    
    # Create table
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Light Grid Accent 1'
    
    # Add header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h
    
    # Add rows
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            if i < len(row_cells):
                row_cells[i].text = cell
    
    return True

def convert_markdown_to_word(md_filepath, docx_filepath):
    """Convert markdown file to Word document"""
    
    # Read markdown
    content = read_markdown_file(md_filepath)
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Split by lines
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Headings
        if line.startswith('# '):
            add_heading_with_level(doc, line, 1)
        elif line.startswith('## '):
            add_heading_with_level(doc, line, 2)
        elif line.startswith('### '):
            add_heading_with_level(doc, line, 3)
        elif line.startswith('#### '):
            add_heading_with_level(doc, line, 4)
        
        # Tables
        elif line.startswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            add_table_from_markdown(doc, '\n'.join(table_lines))
            i -= 1
        
        # Code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            # Light gray background
            from docx.oxml import OxmlElement
            shd = OxmlElement('w:shd')
            shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'E8E8E8')
            p._element.get_or_add_pPr().append(shd)
        
        # Horizontal rule
        elif line.strip() == '---':
            doc.add_paragraph('_' * 80)
        
        # Blockquote
        elif line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
            p.runs[0].italic = True
        
        # Bullet list
        elif line.startswith('- '):
            text = line[2:].strip()
            doc.add_paragraph(text, style='List Bullet')
        
        # Numbered list
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line).strip()
            doc.add_paragraph(text, style='List Number')
        
        # Bold/Italic/Code
        elif line.strip():
            # Process inline formatting
            p = doc.add_paragraph()
            
            # Split by ** and ** for bold
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                else:
                    p.add_run(part)
        
        i += 1
    
    # Save document
    doc.save(docx_filepath)
    print(f"✅ Word document created: {docx_filepath}")

if __name__ == '__main__':
    md_file = 'HELIX_MASTER_REPORT_COMPLIANCE_VERIFICATION.md'
    docx_file = 'HELIX_MASTER_REPORT_COMPLIANCE_VERIFICATION.docx'
    
    print(f"Converting {md_file} to {docx_file}...")
    convert_markdown_to_word(md_file, docx_file)
